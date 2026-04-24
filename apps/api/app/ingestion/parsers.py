from __future__ import annotations

import csv
from dataclasses import dataclass
from io import BytesIO, StringIO
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader


class UnsupportedDocumentTypeError(ValueError):
    pass


@dataclass(slots=True)
class ParsedDocument:
    text: str
    parser_name: str
    extraction_notes: str | None = None


def parse_document_bytes(filename: str, file_bytes: bytes) -> ParsedDocument:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _parse_pdf(file_bytes)
    if suffix == ".docx":
        return _parse_docx(file_bytes)
    if suffix == ".txt":
        return ParsedDocument(
            text=file_bytes.decode("utf-8", errors="ignore"),
            parser_name="text",
        )
    if suffix == ".csv":
        return _parse_csv(file_bytes)
    raise UnsupportedDocumentTypeError(f"Unsupported file type: {suffix}")


def _parse_pdf(file_bytes: bytes) -> ParsedDocument:
    reader = PdfReader(BytesIO(file_bytes))
    pages = [page.extract_text() or "" for page in reader.pages]
    return ParsedDocument(
        text="\n\n".join(page.strip() for page in pages if page.strip()),
        parser_name="pypdf",
        extraction_notes=f"Extracted {len(reader.pages)} pages",
    )


def _parse_docx(file_bytes: bytes) -> ParsedDocument:
    doc = DocxDocument(BytesIO(file_bytes))
    paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    return ParsedDocument(
        text="\n\n".join(paragraphs),
        parser_name="python-docx",
        extraction_notes=f"Extracted {len(paragraphs)} non-empty paragraphs",
    )


def _parse_csv(file_bytes: bytes) -> ParsedDocument:
    decoded = file_bytes.decode("utf-8", errors="ignore")
    reader = csv.DictReader(StringIO(decoded))
    rows = []
    for index, row in enumerate(reader, start=1):
        row_text = "; ".join(f"{header}={value}" for header, value in row.items())
        rows.append(f"Row {index}: {row_text}")
    return ParsedDocument(
        text="\n".join(rows),
        parser_name="csv",
        extraction_notes=f"Extracted {len(rows)} rows",
    )

